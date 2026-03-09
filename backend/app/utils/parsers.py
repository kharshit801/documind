"""File parsers that convert uploaded documents to plain text.

Supported formats: PDF (via PyMuPDF), DOCX (via python-docx), and plain
text/markdown. Each parser returns a single newline-joined string. Raise
``UnsupportedFileTypeError`` for anything we don't know how to handle.
"""

from __future__ import annotations

import io
import logging
from pathlib import PurePath
from typing import Callable, Dict

import fitz  # PyMuPDF
from docx import Document

logger = logging.getLogger(__name__)


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


class UnsupportedFileTypeError(ValueError):
    """Raised when an uploaded file extension is not supported."""


class FileParseError(RuntimeError):
    """Raised when an otherwise-supported file fails to parse."""


def _parse_pdf(content: bytes) -> str:
    try:
        with fitz.open(stream=content, filetype="pdf") as doc:
            pages = [page.get_text("text") for page in doc]
        text = "\n".join(p for p in pages if p)
    except Exception as exc:  # noqa: BLE001 - normalize to FileParseError
        raise FileParseError(f"Failed to parse PDF: {exc}") from exc
    return text.strip()


def _parse_docx(content: bytes) -> str:
    try:
        document = Document(io.BytesIO(content))
        paragraphs = [p.text for p in document.paragraphs if p.text]
        # Pull table cell text too; many real-world DOCX files have content
        # in tables that would otherwise be silently dropped.
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        paragraphs.append(cell.text)
        text = "\n".join(paragraphs)
    except Exception as exc:  # noqa: BLE001
        raise FileParseError(f"Failed to parse DOCX: {exc}") from exc
    return text.strip()


def _parse_text(content: bytes) -> str:
    # Try utf-8 first, fall back to latin-1 to avoid hard-failing on odd files.
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return content.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    raise FileParseError("Unable to decode text file with utf-8 or latin-1")


_PARSERS: Dict[str, Callable[[bytes], str]] = {
    ".pdf": _parse_pdf,
    ".docx": _parse_docx,
    ".txt": _parse_text,
    ".md": _parse_text,
}


def parse_file(filename: str, content: bytes) -> str:
    """Dispatch to the right parser based on file extension.

    Args:
        filename: original filename, used only to determine extension.
        content: raw bytes of the uploaded file.

    Returns:
        Extracted plain text. May be empty if the file has no text.

    Raises:
        UnsupportedFileTypeError: extension is not in ``SUPPORTED_EXTENSIONS``.
        FileParseError: parser failed on a supported extension.
    """
    ext = PurePath(filename).suffix.lower()
    parser = _PARSERS.get(ext)
    if parser is None:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{ext}'. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    text = parser(content)
    if not text:
        logger.warning("Parsed file '%s' produced no text", filename)
    return text
